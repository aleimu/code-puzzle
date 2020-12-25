__doc__ = "mysql数据库生产word文档"

from docx import Document
from pymysql import connect

db = 'mysql_db'
host = '127.0.0.1'
user = 'root'
password = 'passwd'
port = 3306

conn = connect(host=host, user=user, passwd=password, db=db, port=port)
cursor = conn.cursor()
cursor.execute(f"select table_name,table_comment from information_schema.tables where table_schema = '{db}'")
document = Document()
for (table_name, table_comment) in cursor.fetchall():
    document.add_heading(table_name + " " + table_comment, level=1)
    # document.add_paragraph(table_comment)
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'字段'
    hdr_cells[1].text = u'类型'
    hdr_cells[2].text = u'说明'
    cursor.execute(f"SHOW FULL FIELDS FROM {db}.%s" % table_name)
    for (field, type, _, _, _, _, _, _, comment) in cursor.fetchall():
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = type
        row_cells[2].text = comment

document.save('mysql_db.docx')
